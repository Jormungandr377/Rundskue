"""Generate Word (.docx) versions of all policy documents."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_style(doc):
    """Configure base document style."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for i in range(1, 4):
        heading = doc.styles[f'Heading {i}']
        heading.font.name = 'Calibri'
        heading.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)  # teal-600
        if i == 1:
            heading.font.size = Pt(24)
        elif i == 2:
            heading.font.size = Pt(16)
        elif i == 3:
            heading.font.size = Pt(13)


def add_meta(doc, title, version="1.0", date="February 7, 2026"):
    """Add title and metadata block."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RUNDSKUE")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)
    run.font.bold = True

    doc.add_heading(title, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version {version}  |  Effective: {date}  |  finance.rundskue.com").font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x78, 0x78, 0x78)
    doc.add_paragraph()  # spacer


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f" {text}")
    else:
        p.add_run(text)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = val

    doc.add_paragraph()  # spacer


# ============================================================
# SECURITY POLICY
# ============================================================
def create_security_policy():
    doc = Document()
    set_style(doc)
    add_meta(doc, "Security Policy")

    # 1
    doc.add_heading("1. Purpose & Scope", level=2)
    doc.add_paragraph(
        "This document describes the security measures implemented in Finance Tracker, "
        "a personal finance tracking application built and maintained by Rundskue. It provides "
        "transparency about how user data is protected and establishes baseline security practices."
    )
    doc.add_paragraph(
        "This policy applies to the Finance Tracker web application at finance.rundskue.com, "
        "including the FastAPI backend, React frontend, PostgreSQL database, Docker infrastructure, "
        "and third-party integrations (Plaid API, Sentry)."
    )
    p = doc.add_paragraph()
    run = p.add_run(
        "Disclaimer: Finance Tracker is a small, self-hosted personal project. While reasonable "
        "security measures have been implemented, this application has not undergone a formal "
        "third-party security audit and is not SOC 2 or ISO 27001 certified."
    )
    run.italic = True
    run.font.size = Pt(10)

    # 2
    doc.add_heading("2. Data Classification", level=2)
    doc.add_heading("Critical", level=3)
    add_bullet(doc, "Encrypted at rest with Fernet symmetric encryption. Never logged or exposed in API responses.", "Plaid access tokens:")
    add_bullet(doc, "Hashed with bcrypt before storage. Plaintext passwords are never stored or logged.", "User passwords:")
    add_bullet(doc, "Stored as environment variables. Never committed to source control.", "JWT signing secrets:")

    doc.add_heading("Sensitive", level=3)
    add_bullet(doc, "Financial transaction data, account balances, and account metadata")
    add_bullet(doc, "User profile information (email addresses, display names)")
    add_bullet(doc, "Session and device data, refresh tokens")
    add_bullet(doc, "Two-factor authentication TOTP secrets")

    doc.add_heading("Internal", level=3)
    add_bullet(doc, "Application logs and error traces (sent to Sentry)")
    add_bullet(doc, "Non-secret application configuration and deployment settings")

    # 3
    doc.add_heading("3. Access Control", level=2)
    add_bullet(doc, "All users must register and authenticate before accessing any financial data.")
    add_bullet(doc, "API endpoints enforce ownership checks \u2014 users can only access their own data.")
    add_bullet(doc, "An admin role exists for elevated privileges with no self-service escalation.")
    add_bullet(doc, "CORS whitelist permits requests only from finance.rundskue.com and localhost (development).")
    add_bullet(doc, "PostgreSQL database is not exposed to the public internet \u2014 accessible only within the Docker network.")
    add_bullet(doc, "Database credentials stored as environment variables, never in source control.")

    # 4
    doc.add_heading("4. Encryption", level=2)
    doc.add_heading("In Transit", level=3)
    add_bullet(doc, "All traffic served over HTTPS with TLS encryption.")
    add_bullet(doc, "TLS certificates automatically managed and renewed.")
    add_bullet(doc, "Unencrypted HTTP requests redirected to HTTPS.")

    doc.add_heading("At Rest", level=3)
    add_bullet(doc, "Plaid access tokens encrypted using Fernet symmetric encryption before database storage.")
    add_bullet(doc, "Passwords hashed using bcrypt (one-way hash, cannot be decrypted).")
    p = doc.add_paragraph()
    run = p.add_run(
        "Transparency note: Financial transaction data and user profiles are stored in PostgreSQL "
        "without field-level encryption. Protection relies on network isolation and server access controls."
    )
    run.italic = True
    run.font.size = Pt(10)

    # 5
    doc.add_heading("5. Authentication & Session Management", level=2)
    add_bullet(doc, "Short-lived access tokens and longer-lived refresh tokens.", "JWT-based auth:")
    add_bullet(doc, "TOTP-based, compatible with Google Authenticator, Authy, and similar apps.", "Two-factor authentication:")
    add_bullet(doc, "Users can view and revoke active sessions.", "Session & device tracking:")
    add_bullet(doc, "Refresh tokens can be revoked on logout or password change.")
    add_bullet(doc, "Passwords validated via Pydantic models on the backend.")

    # 6
    doc.add_heading("6. API Security", level=2)
    add_bullet(doc, "X-Requested-With header validation on mutating requests.", "CSRF protection:")
    add_bullet(doc, "slowapi rate limiting to prevent abuse and brute-force attacks.", "Rate limiting:")
    add_bullet(doc, "Pydantic models reject malformed input before it reaches business logic.", "Input validation:")
    add_bullet(doc, "SQLAlchemy ORM parameterized queries.", "SQL injection prevention:")
    add_bullet(doc, "Middleware for optimized response delivery.", "GZip compression:")

    # 7
    doc.add_heading("7. Infrastructure Security", level=2)
    add_bullet(doc, "Deployed as Docker containers on a self-hosted server, managed through Coolify.")
    add_bullet(doc, "Backend, frontend, and database run in isolated containers within a shared Docker network.")
    add_bullet(doc, "Only ports 80/443 exposed to the public internet.")
    add_bullet(doc, "All secrets stored as environment variables, never in source code.")
    add_bullet(doc, "Runtime errors monitored via Sentry (configured to minimize sensitive data exposure).")
    p = doc.add_paragraph()
    run = p.add_run(
        "Known limitations: Single maintainer, no dedicated security operations team, no IDS, "
        "no automated vulnerability scanning, and no WAF beyond the reverse proxy. "
        "OS and Docker updates applied on a best-effort basis."
    )
    run.italic = True
    run.font.size = Pt(10)

    # 8
    doc.add_heading("8. Incident Response", level=2)
    doc.add_paragraph("In the event of a suspected security incident:")
    add_bullet(doc, "Determine nature and scope using available logs and Sentry data.", "Assess:")
    add_bullet(doc, "Revoke compromised tokens, rotate secrets, disable affected accounts, or take the application offline.", "Contain:")
    add_bullet(doc, "Inform affected users promptly with an honest description of the incident.", "Notify:")
    add_bullet(doc, "Fix the underlying vulnerability, rotate credentials, restore operations.", "Remediate:")
    add_bullet(doc, "Document what happened and prevent recurrence.", "Review:")
    doc.add_paragraph(
        "If Plaid access tokens are suspected compromised, the Fernet encryption key is rotated "
        "immediately, all tokens invalidated via Plaid API, and users asked to re-link accounts."
    )

    # 9
    doc.add_heading("9. Third-Party Integrations", level=2)
    doc.add_heading("Plaid", level=3)
    doc.add_paragraph(
        "Users connect bank accounts through Plaid Link. Credentials are entered directly into "
        "Plaid's interface and are never seen or stored by Finance Tracker. Plaid returns an "
        "access token (encrypted before storage) used to fetch account and transaction data. "
        "Plaid is SOC 2 Type II certified with AES-256 encryption. See plaid.com/security/ for details."
    )
    doc.add_heading("Sentry", level=3)
    doc.add_paragraph(
        "Used for error monitoring. Reports may include technical metadata but are configured "
        "to minimize sensitive data exposure. See sentry.io/security/ for details."
    )

    # 10
    doc.add_heading("10. Contact Information", level=2)
    doc.add_paragraph("For security concerns, vulnerability reports, or questions about this policy:")
    add_table(doc, ["Channel", "Contact"], [
        ["Maintainer", "Luke Robinson"],
        ["Email", "rundskue@outlook.com"],
        ["Application", "https://finance.rundskue.com"],
    ])
    doc.add_paragraph(
        "If you discover a security vulnerability, please report it privately rather than "
        "opening a public issue. Responsible disclosure is appreciated."
    )

    path = os.path.join(OUTPUT_DIR, "Security_Policy_Rundskue.docx")
    doc.save(path)
    print(f"Created: {path}")


# ============================================================
# PRIVACY POLICY
# ============================================================
def create_privacy_policy():
    doc = Document()
    set_style(doc)
    add_meta(doc, "Privacy Policy")

    # 1
    doc.add_heading("1. Introduction", level=2)
    doc.add_paragraph(
        "Rundskue operates Finance Tracker, a personal financial management tool available at "
        "finance.rundskue.com. This Privacy Policy explains how we collect, use, and protect "
        "your information when you use our service."
    )

    # 2
    doc.add_heading("2. Information We Collect", level=2)
    doc.add_heading("Account Information", level=3)
    doc.add_paragraph(
        "When you create an account, we collect your email address and password. "
        "Your password is hashed using bcrypt and is never stored in plain text."
    )
    doc.add_heading("Financial Data", level=3)
    doc.add_paragraph(
        "If you choose to link a bank account, we receive bank account details and transaction "
        "data synced through Plaid. This data is used solely to provide you with financial "
        "tracking and budgeting features."
    )
    doc.add_heading("Usage Data", level=3)
    doc.add_paragraph(
        "We collect session information, device type, and login timestamps to manage your "
        "active sessions and ensure account security."
    )

    # 3
    doc.add_heading("3. How We Use Your Information", level=2)
    add_bullet(doc, "To provide financial tracking, budgeting, and reporting features.")
    add_bullet(doc, "To sync your bank accounts and transactions via Plaid.")
    add_bullet(doc, "To send notifications about your accounts (budget alerts, bill reminders).")
    add_bullet(doc, "For error monitoring via Sentry (no PII is transmitted to Sentry).")

    # 4
    doc.add_heading("4. Third-Party Services", level=2)
    doc.add_heading("Plaid", level=3)
    doc.add_paragraph(
        "We use Plaid to securely connect your bank accounts. When you link an account, "
        "Plaid collects and processes your banking credentials according to their own privacy "
        "policy. We encourage you to review Plaid's privacy policy at plaid.com/legal/."
    )
    doc.add_heading("Sentry", level=3)
    doc.add_paragraph(
        "We use Sentry for error monitoring and application performance tracking. No personal "
        "financial data or personally identifiable information is shared with Sentry."
    )

    # 5
    doc.add_heading("5. Data Security", level=2)
    add_bullet(doc, "All data in transit is protected with HTTPS/TLS encryption.")
    add_bullet(doc, "Plaid access tokens are encrypted at rest using Fernet symmetric encryption.")
    add_bullet(doc, "Passwords are hashed with bcrypt and never stored in plain text.")
    add_bullet(doc, "Two-factor authentication (2FA) is available via TOTP-based authenticator apps.")

    # 6
    doc.add_heading("6. Data Retention", level=2)
    doc.add_paragraph(
        "Your financial data is retained for as long as your account remains active. "
        "If you wish to have your data deleted, you may request account deletion at any time. "
        "All associated data will be permanently removed within 30 days."
    )

    # 7
    doc.add_heading("7. Your Rights", level=2)
    add_bullet(doc, "You can view all stored data within the application at any time.", "Access:")
    add_bullet(doc, "You can export your transaction data in CSV or Excel format.", "Export:")
    add_bullet(doc, "You can request complete deletion of your account and all associated data.", "Deletion:")

    # 8
    doc.add_heading("8. Contact", level=2)
    doc.add_paragraph("For questions or data rights requests:")
    add_table(doc, ["", ""], [
        ["Name", "Luke Robinson"],
        ["Email", "rundskue@outlook.com"],
    ])

    # 9
    doc.add_heading("9. Changes to This Policy", level=2)
    doc.add_paragraph(
        "We may update this Privacy Policy from time to time. Changes will be reflected on "
        "this page with an updated effective date. Last updated: February 2026."
    )

    path = os.path.join(OUTPUT_DIR, "Privacy_Policy_Rundskue.docx")
    doc.save(path)
    print(f"Created: {path}")


# ============================================================
# DATA RETENTION POLICY
# ============================================================
def create_data_retention_policy():
    doc = Document()
    set_style(doc)
    add_meta(doc, "Data Retention Policy")

    # 1
    doc.add_heading("1. Purpose & Scope", level=2)
    doc.add_paragraph(
        "This document defines the data retention practices for Finance Tracker, a personal "
        "finance management application operated by Rundskue at finance.rundskue.com. It describes "
        "what data is collected, how long it is retained, and how it is deleted."
    )

    # 2
    doc.add_heading("2. Data Categories & Retention Periods", level=2)

    doc.add_heading("User Account Data", level=3)
    add_table(doc, ["Data", "Retention Period"], [
        ["Email address", "While account is active; deleted within 30 days of deletion request"],
        ["Hashed password", "While account is active; deleted within 30 days of deletion request"],
        ["2FA secrets", "While 2FA is enabled; deleted when disabled or account deleted"],
    ])

    doc.add_heading("Bank Connection Data (Plaid)", level=3)
    add_table(doc, ["Data", "Retention Period"], [
        ["Plaid access tokens (encrypted)", "While bank link is active; revoked via Plaid API and deleted on unlink"],
        ["Account info (name, type, last 4, balances)", "While bank link is active; deleted on unlink or account deletion"],
    ])

    doc.add_heading("Financial Data", level=3)
    add_table(doc, ["Data", "Retention Period"], [
        ["Transactions", "While account is active; deleted within 30 days of deletion request"],
        ["Budget configurations", "While account is active; deleted within 30 days of deletion request"],
        ["Savings goals", "While account is active; deleted within 30 days of deletion request"],
        ["Recurring bills", "While account is active; deleted within 30 days of deletion request"],
    ])

    doc.add_heading("Session & Log Data", level=3)
    add_table(doc, ["Data", "Retention Period"], [
        ["Session tokens & device info", "Auto-expired after 30 days of inactivity"],
        ["Application logs", "Retained for 90 days, then automatically deleted"],
    ])

    # 3
    doc.add_heading("3. Data Deletion Procedures", level=2)

    doc.add_heading("Account Deletion", level=3)
    doc.add_paragraph("When a user requests account deletion:")
    add_bullet(doc, "Account is deactivated immediately (login disabled).")
    add_bullet(doc, "All Plaid access tokens are revoked via the Plaid API, then deleted.")
    add_bullet(doc, "All user data is permanently deleted from the database within 30 days.")
    add_bullet(doc, "User is notified by email once deletion is complete.")

    doc.add_heading("Bank Account Unlinking", level=3)
    add_bullet(doc, "Plaid access token is revoked via the Plaid API.")
    add_bullet(doc, "Token and bank account metadata are deleted from the database.")
    add_bullet(doc, "Previously synced transactions are retained unless user explicitly requests deletion.")

    doc.add_heading("Automated Expiration", level=3)
    add_bullet(doc, "Sessions inactive for over 30 days are automatically purged.")
    add_bullet(doc, "Application logs older than 90 days are automatically removed.")

    p = doc.add_paragraph()
    run = p.add_run(
        "Note: Data is deleted via standard SQL DELETE operations. Deleted data may persist "
        "in database backups for up to 30 days before aging out."
    )
    run.italic = True
    run.font.size = Pt(10)

    # 4
    doc.add_heading("4. Your Rights", level=2)
    add_bullet(doc, "Request deletion of your account and all associated data at any time. Completed within 30 days.", "Deletion:")
    add_bullet(doc, "Export your transaction data in CSV or Excel format through the application.", "Export:")
    add_bullet(doc, "View all your stored data within the application at any time.", "Access:")

    # 5
    doc.add_heading("5. Backup & Archive", level=2)
    add_bullet(doc, "PostgreSQL database is backed up daily.")
    add_bullet(doc, "Backups retained for up to 30 days; older backups automatically removed.")
    add_bullet(doc, "Deleted data may persist in backups for up to 30 days after live deletion.")
    add_bullet(doc, "No separate long-term archive is maintained.")

    # 6
    doc.add_heading("6. Review Schedule", level=2)
    doc.add_paragraph(
        "This policy is reviewed at least once per year. Next scheduled review: February 2027. "
        "Significant changes will be communicated to users."
    )

    # 7
    doc.add_heading("7. Contact", level=2)
    doc.add_paragraph("For questions or data-related requests:")
    add_table(doc, ["", ""], [
        ["Name", "Luke Robinson"],
        ["Email", "rundskue@outlook.com"],
    ])

    path = os.path.join(OUTPUT_DIR, "Data_Retention_Policy_Rundskue.docx")
    doc.save(path)
    print(f"Created: {path}")


if __name__ == "__main__":
    create_security_policy()
    create_privacy_policy()
    create_data_retention_policy()
    print("\nAll documents generated successfully!")
